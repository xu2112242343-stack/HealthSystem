#!/usr/bin/env python3
"""生成《三病传播展示分计算说明》Word 文档（数据密度动态权重版）。"""

from __future__ import annotations

from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn
except ImportError as exc:
    raise SystemExit("请先安装: pip install python-docx") from exc


OUT_PATH = Path(__file__).resolve().parent / "三病传播展示分计算说明.docx"


def set_doc_default_font(doc: Document, name: str = "宋体", size_pt: int = 12) -> None:
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size_pt)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_formula_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = val


def build() -> Document:
    doc = Document()
    set_doc_default_font(doc)
    today = date.today().isoformat()

    add_title(doc, "三病传播展示分计算说明")
    add_p(doc, f"生成日期：{today}")
    add_p(doc, "代码位置：web/backend/app/risk_engine.py → propagation_scores_fusion")
    doc.add_paragraph()

    add_h(doc, "一、设计说明")
    add_bullets(
        doc,
        [
            "在保留原算法整体结构的前提下，将 dis_core 中固定的 0.62、0.38 改为按数据密度动态计算。",
            "不引入关联-因果双层、混淆分解等复杂模块，便于理解与答辩说明。",
            "输出为三角图三条边的传播展示分（0~98），非单病概率，也不替代六模型结果。",
        ],
    )

    add_h(doc, "二、输入与输出")
    add_table(
        doc,
        ["输入", "说明"],
        [
            ["p_L, p_D, p_S", "三病概率 [0,1]"],
            ["各病 Top 因子", "名称 + 展示权重 value"],
        ],
    )
    doc.add_paragraph()
    add_table(
        doc,
        ["输出", "说明"],
        [
            ["propagationScores[3]", "顺序：糖尿病→脂肪肝、脂肪肝→脑卒中、糖尿病→脑卒中"],
            ["propagationDetail", "含 wPair、wG3、disCore、诊断文案等"],
        ],
    )

    add_h(doc, "三、三条有向边")
    add_table(
        doc,
        ["边", "p_src", "p_tgt", "因子向量"],
        [
            ["糖尿病→脂肪肝", "p_D", "p_L", "wd, wl"],
            ["脂肪肝→脑卒中", "p_L", "p_S", "wl, ws"],
            ["糖尿病→脑卒中", "p_D", "p_S", "wd, ws"],
        ],
    )

    add_h(doc, "四、计算公式（逐步）")

    add_h(doc, "4.1 符号说明", level=2)
    add_p(doc, "下列符号对三条有向边通用；计算某一条边时，p_src、p_tgt 及因子向量按第三节对应关系取值。")
    add_table(
        doc,
        ["符号", "含义", "取值范围 / 说明"],
        [
            ["p_L", "脂肪肝（MAFLD）风险概率", "[0, 1]，由肝病模型/融合得到"],
            ["p_D", "糖尿病（T2DM）风险概率", "[0, 1]，由糖尿病模型/融合得到"],
            ["p_S", "脑卒中（CVA）风险概率", "[0, 1]，由卒中模型/融合得到"],
            ["p_src", "当前有向边「源病」概率", "糖→肝取 p_D；肝→卒取 p_L；糖→卒取 p_D"],
            ["p_tgt", "当前有向边「靶病」概率", "糖→肝取 p_L；肝→卒取 p_S；糖→卒取 p_S"],
            ["pair_geo", "该边两端概率的几何平均（两病共现强度）", "[0, 1]，见步骤 1"],
            ["G₃", "三病概率的几何平均（三病整体共病背景）", "[0, 1]，见步骤 2"],
            ["d_pair", "pair_geo 的数据密度", "等于 pair_geo"],
            ["d_g3", "G₃ 的数据密度", "等于 G₃"],
            ["w_pair", "pair_geo 在 dis_core 中的动态权重", "[0, 1]，且 w_pair + w_g3 = 1"],
            ["w_g3", "G₃ 在 dis_core 中的动态权重", "[0, 1]，见步骤 4"],
            ["dis_core", "疾病层合成强度", "[0, 1]，见步骤 5"],
            ["cos_AB", "边两端 Top 因子权重向量的余弦相似度", "[0, 1]，无因子时为 0"],
            ["blend", "疾病层与因子层合成后的最终强度", "[0, 1]，见步骤 7"],
            ["impact", "三角图弧上传播展示分", "整数，0~98，见步骤 8"],
            ["wl / wd / ws", "脂肪肝 / 糖尿病 / 脑卒中 Top 因子权重字典", "键为归一化因子名，值为展示权重"],
        ],
    )
    doc.add_paragraph()

    add_h(doc, "4.2 逐步计算", level=2)
    steps = [
        (
            "步骤 1：计算该边两端几何耦合 pair_geo",
            "pair_geo = √(p_src × p_tgt)",
            "p_src、p_tgt 为当前边的源、靶病概率；几何平均反映两病同时处于较高风险域的强度。",
        ),
        (
            "步骤 2：计算三病整体几何平均 G₃",
            "G₃ = (p_L × p_D × p_S)^(1/3)",
            "三个指数 1/3 表示开三次方；G₃ 反映三病整体代谢/血管风险共性背景。",
        ),
        (
            "步骤 3：定义数据密度",
            "d_pair = pair_geo，d_g3 = G₃",
            "密度即用于分配权重的强度值；两端共现用 d_pair，三病整体用 d_g3。",
        ),
        (
            "步骤 4：按密度计算动态权重（替代固定 0.62 / 0.38）",
            "w_pair = d_pair / (d_pair + d_g3)\nw_g3 = d_g3 / (d_pair + d_g3)",
            "分母为两密度之和；w_pair 越大表示该边两端共现在 dis_core 中占比越高。",
        ),
        (
            "步骤 5：合成疾病层强度 dis_core",
            "dis_core = w_pair × pair_geo + w_g3 × G₃",
            "原算法为 0.62×pair_geo + 0.38×G₃；现改为 w_pair、w_g3 随患者数据变化。",
        ),
        (
            "步骤 6：计算因子余弦 cos_AB",
            "cos_AB = cos(向量 A, 向量 B)",
            "A、B 为边两端 Top 因子权重向量（同名因子对齐）；衡量重要因子的共性，无因子则 cos_AB=0。",
        ),
        (
            "步骤 7：合成 blend",
            "有因子：blend = 0.5 × dis_core + 0.5 × cos_AB\n无因子：blend = dis_core",
            "0.5 为疾病层与因子层的固定混合比例，与原版一致。",
        ),
        (
            "步骤 8：得到传播展示分 impact",
            "impact = round(100 × blend)，并限制在 [0, 98]",
            "round 为四舍五入取整；前端弧上显示为「传播↑impact%」。",
        ),
    ]
    for i, (title, formula, note) in enumerate(steps, 1):
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        add_formula_block(doc, formula.split("\n"))
        add_p(doc, f"说明：{note}")
        if i < len(steps):
            doc.add_paragraph()

    add_h(doc, "五、动态权重与固定权重的关系")
    add_p(
        doc,
        "当 pair_geo 与 G₃ 接近时，w_pair≈0.5、w_g3≈0.5，结果接近二者算术平均。"
        "当 pair_geo 明显大于 G₃ 时，w_pair 升高（可高于原 0.62），更突出该边两端共现；"
        "当 G₃ 明显大于 pair_geo 时，w_g3 升高（可高于原 0.38），更突出三病整体代谢/血管背景。"
    )

    add_h(doc, "六、与原算法对照")
    add_table(
        doc,
        ["项目", "原算法", "现算法"],
        [
            ["dis_core 权重", "固定 0.62 / 0.38", "按 pair_geo、G₃ 密度比例动态分配"],
            ["因子融合", "0.5·dis_core + 0.5·cos", "不变"],
            ["最终展示分", "round(100·blend)", "不变"],
        ],
    )

    add_h(doc, "七、数值示例（糖→肝边）")
    add_p(doc, "设 p_D=0.62，p_L=0.55，p_S=0.48，且无有效因子（cos_AB=0）：")
    add_formula_block(
        doc,
        [
            "pair_geo = sqrt(0.62 x 0.55) = 0.584",
            "G3 = (0.62 x 0.55 x 0.48)^(1/3) = 0.547",
            "w_pair = 0.584 / (0.584 + 0.547) = 0.516  ->  约 52%",
            "w_g3   = 0.547 / (0.584 + 0.547) = 0.484  ->  约 48%",
            "dis_core = 0.516 x 0.584 + 0.484 x 0.547 = 0.566",
            "blend = dis_core = 0.566  ->  impact = 57",
        ],
    )
    add_p(doc, "若原固定权重 0.62/0.38，则 dis_core = 0.62x0.584 + 0.38x0.547 = 0.570，略有差异。")

    add_h(doc, "八、propagationDetail 主要字段")
    add_table(
        doc,
        ["字段", "含义"],
        [
            ["impact", "弧上展示分（0~98）"],
            ["decomposition.wPair / wG3", "数据密度动态权重"],
            ["decomposition.densityPair / densityG3", "密度值 pair_geo 与 G3"],
            ["decomposition.disCore", "疾病层合成强度"],
            ["decomposition.cosAb", "因子余弦相似度"],
            ["diagnosis.label", "数据密度加权说明文案"],
        ],
    )

    add_h(doc, "九、API 与部署")
    add_bullets(
        doc,
        [
            "接口：POST /api/risk/predict",
            "修改后需重启后端进程方可生效。",
        ],
    )

    return doc


def main() -> None:
    doc = build()
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"已生成: {OUT_PATH}")


if __name__ == "__main__":
    main()
