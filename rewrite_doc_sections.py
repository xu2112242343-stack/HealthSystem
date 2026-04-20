from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import os

base = r"c:\Users\21122\xwechat_files\wxid_p6jnc01rxek322_9c5c\msg\file\2026-04"
files = [f for f in os.listdir(base) if f.lower().endswith(".docx") and f != "s202308384144-ex02.docx"]
path = os.path.join(base, files[0])
out_path = r"c:\Users\21122\Desktop\HealthSystem\chapter5_6_real_rewrite.docx"

doc = Document(path)


def find_idx(startswith: str):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(startswith):
            return i
    return None


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def insert_after(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    out = Paragraph(new_p, paragraph._parent)
    out.add_run(text)
    return out


sec5 = find_idx("第5章")
sec6 = find_idx("第6章")
sec7 = find_idx("第7章")
if sec5 is None or sec6 is None or sec7 is None:
    raise SystemExit("未找到第5/6/7章标题，未修改。")

for i in range(sec6 - 1, sec5, -1):
    delete_paragraph(doc.paragraphs[i])

p = doc.paragraphs[sec5]
sec5_lines = [
    "5.1 系统工程架构与技术栈",
    "本项目采用前后端分离架构。前端为 React + TypeScript + Tailwind CSS（Vite 构建），后端为 FastAPI + SQLAlchemy + MySQL。核心能力围绕“肝病、糖尿病、脑卒中”三病风险评估与健康干预展开。",
    "后端以 REST API 提供用户、问卷、风险预测、健康指南、医院推荐、图像上传与读取等服务；前端实现用户端与管理端的业务流程页面，完成数据采集、风险展示、推荐解释与干预交互。",
    "5.2 后端核心实现（FastAPI）",
    "（1）统一预测入口：后端通过 risk_engine.predict_triple 计算三病风险，输出概率、风险等级、分数与可解释因素。风险分层阈值按当前项目配置：<30% 低风险、[30%,60%) 中风险、≥60% 高风险。",
    "（2）结构化与图像融合：针对三病分别支持结构化数据概率与图像概率融合，最终概率采用结构化:图像=6:4 的加权策略；若仅存在单一路径数据，则直接采用该路概率。",
    "（3）健康指南推荐：推荐逻辑采用规则引擎而非纯排序打分。三病均低风险时，优先推送“认知类+饮食/运动类”；存在中高风险时，按病种与风险等级精准匹配，并对标题含病种关键词的文章前置。",
    "（4）医院推荐：基于前端定位坐标与医院经纬度，前端按 Haversine 距离计算并排序，实现“就近就医推荐”。",
    "5.3 图像数据链路实现（上传-落盘-入库-预览-删除）",
    "项目在 user_info 表新增 liver_image_path、diabetes_image_path、stroke_image_path 三个字段。图像上传后由后端落盘并把绝对路径写入对应字段；提供元数据接口与文件读取接口用于前端预览；删除时同步清理数据库路径与本地文件。",
    "该实现已支持“再次进入页面可回显已上传图像、可覆盖修改、可删除重传”，满足持续管理场景。",
    "5.4 前端核心实现（React）",
    "（1）风险与干预页面联动：风险页展示三病概率、风险等级与综合分；干预页展示“疾病风险程度”、个性化指南与医院推荐。",
    "（2）健康指南弹窗：点击文章卡片可打开全文弹窗，支持正文分段与图片占位符替换，提升可读性与可解释性。",
    "（3）数据采集与上传交互：DataCollection 页面完成问卷采集与分病种图像上传；上传组件支持本地文件预览、远端已上传内容回显、下载/打开、删除等完整交互。",
    "5.5 数据与脚本工程化",
    "项目提供脚本完成外部 CSV 到 user_info 字段映射、批量创建用户与健康数据初始化，便于演示环境快速构建与复现实验。数据库初始化阶段包含 user_info 新增字段的自动检查与补齐逻辑，降低部署门槛。",
    "5.6 部署与运行方式（当前项目真实状态）",
    "当前版本以本地开发/演示部署为主：后端服务、前端服务与数据库可在同一环境启动。工程中已具备模块化接口、脚本化数据初始化和可持续扩展的服务边界，后续可平滑扩展到容器化与多节点部署。",
]
for line in sec5_lines:
    p = insert_after(p, line)

sec6 = find_idx("第6章")
sec7 = find_idx("第7章")
if sec6 is None or sec7 is None:
    raise SystemExit("重定位第6/7章失败，未继续修改第6章。")

for i in range(sec7 - 1, sec6, -1):
    delete_paragraph(doc.paragraphs[i])

p = doc.paragraphs[sec6]
sec6_lines = [
    "6.1 测试范围与原则",
    "本项目测试覆盖“模型预测链路、前后端核心业务链路、图像上传与回显链路、推荐逻辑链路”。测试遵循“功能可用、逻辑正确、结果可解释、回归不破坏”四项原则，避免仅展示静态界面。",
    "6.2 风险预测与融合逻辑验证",
    "（1）结构化预测验证：使用问卷数据驱动三病风险预测，检查概率输出、风险分层与因素解释是否一致。",
    "（2）融合规则验证：重点验证“结构化:图像=6:4”规则。分别覆盖三类场景：双路都存在、仅结构化存在、仅图像存在，确保最终概率符合业务定义。",
    "（3）阈值回归验证：针对30%与60%分界值做边界检查，保证低/中/高风险标签判定稳定。",
    "6.3 个性化推荐与干预功能验证",
    "（1）健康指南推荐：验证“全低风险推认知+饮食运动；中高风险按病种+等级精准匹配；标题关键词前置”三条核心规则。",
    "（2）文章阅读体验：验证点击文章卡片后弹窗可正常展示完整正文，段落切分与图片替换生效。",
    "（3）医院推荐：验证定位授权、距离计算、按距离排序与异常状态提示（未授权/失败）的处理。",
    "6.4 图像上传闭环验证",
    "（1）上传后入库：验证文件落盘成功，user_info 对应 image_path 字段写入正确。",
    "（2）再次进入回显：验证已上传图像可通过元数据与文件接口回显。",
    "（3）修改与删除：验证覆盖上传后路径更新、删除后数据库清空与文件移除。",
    "（4）稳定性修复回归：针对历史“影像上传点击后白屏”问题完成回归验证，确保页面可正常进入与操作。",
    "6.5 工程质量与运行验证",
    "前端执行 TypeScript 静态检查（如 npx tsc --noEmit）与关键页面流程自测；后端对核心模块进行语法与接口级验证，确保改动后服务可启动、接口响应结构稳定。",
    "6.6 测试结论（基于当前版本）",
    "当前项目已实现从“数据采集-风险预测-图像融合-个性化推荐-干预展示”的端到端闭环。核心业务逻辑、图像上传全链路与推荐策略均可运行，满足课程与竞赛场景下的系统演示和持续迭代需求。后续将进一步补充自动化测试与容器化性能压测，提升工程化成熟度。",
]
for line in sec6_lines:
    p = insert_after(p, line)

doc.save(out_path)
print("SAVED:", out_path)
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import os

base = r'c:\Users\21122\xwechat_files\wxid_p6jnc01rxek322_9c5c\msg\file\2026-04'
files = [f for f in os.listdir(base) if f.lower().endswith('.docx') and f != 's202308384144-ex02.docx']
path = os.path.join(base, files[0])

doc = Document(path)

def find_idx(startswith: str):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(startswith):
            return i
    return None

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def insert_after(paragraph, text: str, style=None):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    out = Paragraph(new_p, paragraph._parent)
    if style is not None:
        out.style = style
    out.add_run(text)
    return out

sec5 = find_idx('第5章')
sec6 = find_idx('第6章')
sec7 = find_idx('第7章')
if sec5 is None or sec6 is None or sec7 is None:
    raise SystemExit('未找到第5/6/7章标题，未修改。')

for i in range(sec6 - 1, sec5, -1):
    delete_paragraph(doc.paragraphs[i])

p = doc.paragraphs[sec5]
sec5_lines = [
'5.1 系统工程架构与技术栈',
'本项目采用前后端分离架构。前端为 React + TypeScript + Tailwind CSS（Vite 构建），后端为 FastAPI + SQLAlchemy + MySQL。核心能力围绕“肝病、糖尿病、脑卒中”三病风险评估与健康干预展开。',
'后端以 REST API 提供用户、问卷、风险预测、健康指南、医院推荐、图像上传与读取等服务；前端实现用户端与管理端的业务流程页面，完成数据采集、风险展示、推荐解释与干预交互。',
'5.2 后端核心实现（FastAPI）',
'（1）统一预测入口：后端通过 risk_engine.predict_triple 计算三病风险，输出概率、风险等级、分数与可解释因素。风险分层阈值按当前项目配置：<30% 低风险、[30%,60%) 中风险、≥60% 高风险。',
'（2）结构化与图像融合：针对三病分别支持结构化数据概率与图像概率融合，最终概率采用结构化:图像=6:4 的加权策略；若仅存在单一路径数据，则直接采用该路概率。',
'（3）健康指南推荐：推荐逻辑采用规则引擎而非纯排序打分。三病均低风险时，优先推送“认知类+饮食/运动类”；存在中高风险时，按病种与风险等级精准匹配，并对标题含病种关键词的文章前置。',
'（4）医院推荐：基于前端定位坐标与医院经纬度，前端按 Haversine 距离计算并排序，实现“就近就医推荐”。',
'5.3 图像数据链路实现（上传-落盘-入库-预览-删除）',
'项目在 user_info 表新增 liver_image_path、diabetes_image_path、stroke_image_path 三个字段。图像上传后由后端落盘并把绝对路径写入对应字段；提供元数据接口与文件读取接口用于前端预览；删除时同步清理数据库路径与本地文件。',
'该实现已支持“再次进入页面可回显已上传图像、可覆盖修改、可删除重传”，满足持续管理场景。',
'5.4 前端核心实现（React）',
'（1）风险与干预页面联动：风险页展示三病概率、风险等级与综合分；干预页展示“疾病风险程度”、个性化指南与医院推荐。',
'（2）健康指南弹窗：点击文章卡片可打开全文弹窗，支持正文分段与图片占位符替换，提升可读性与可解释性。',
'（3）数据采集与上传交互：DataCollection 页面完成问卷采集与分病种图像上传；上传组件支持本地文件预览、远端已上传内容回显、下载/打开、删除等完整交互。',
'5.5 数据与脚本工程化',
'项目提供脚本完成外部 CSV 到 user_info 字段映射、批量创建用户与健康数据初始化，便于演示环境快速构建与复现实验。数据库初始化阶段包含 user_info 新增字段的自动检查与补齐逻辑，降低部署门槛。',
'5.6 部署与运行方式（当前项目真实状态）',
'当前版本以本地开发/演示部署为主：后端服务、前端服务与数据库可在同一环境启动。工程中已具备模块化接口、脚本化数据初始化和可持续扩展的服务边界，后续可平滑扩展到容器化与多节点部署。'
]
for line in sec5_lines:
    p = insert_after(p, line)

sec6 = find_idx('第6章')
sec7 = find_idx('第7章')
if sec6 is None or sec7 is None:
    raise SystemExit('重定位第6/7章失败，未继续修改第6章。')

for i in range(sec7 - 1, sec6, -1):
    delete_paragraph(doc.paragraphs[i])

p = doc.paragraphs[sec6]
sec6_lines = [
'6.1 测试范围与原则',
'本项目测试覆盖“模型预测链路、前后端核心业务链路、图像上传与回显链路、推荐逻辑链路”。测试遵循“功能可用、逻辑正确、结果可解释、回归不破坏”四项原则，避免仅展示静态界面。',
'6.2 风险预测与融合逻辑验证',
'（1）结构化预测验证：使用问卷数据驱动三病风险预测，检查概率输出、风险分层与因素解释是否一致。',
'（2）融合规则验证：重点验证“结构化:图像=6:4”规则。分别覆盖三类场景：双路都存在、仅结构化存在、仅图像存在，确保最终概率符合业务定义。',
'（3）阈值回归验证：针对30%与60%分界值做边界检查，保证低/中/高风险标签判定稳定。',
'6.3 个性化推荐与干预功能验证',
'（1）健康指南推荐：验证“全低风险推认知+饮食运动；中高风险按病种+等级精准匹配；标题关键词前置”三条核心规则。',
'（2）文章阅读体验：验证点击文章卡片后弹窗可正常展示完整正文，段落切分与图片替换生效。',
'（3）医院推荐：验证定位授权、距离计算、按距离排序与异常状态提示（未授权/失败）的处理。',
'6.4 图像上传闭环验证',
'（1）上传后入库：验证文件落盘成功，user_info 对应 image_path 字段写入正确。',
'（2）再次进入回显：验证已上传图像可通过元数据与文件接口回显。',
'（3）修改与删除：验证覆盖上传后路径更新、删除后数据库清空与文件移除。',
'（4）稳定性修复回归：针对历史“影像上传点击后白屏”问题完成回归验证，确保页面可正常进入与操作。',
'6.5 工程质量与运行验证',
'前端执行 TypeScript 静态检查（如 npx tsc --noEmit）与关键页面流程自测；后端对核心模块进行语法与接口级验证，确保改动后服务可启动、接口响应结构稳定。',
'6.6 测试结论（基于当前版本）',
'当前项目已实现从“数据采集-风险预测-图像融合-个性化推荐-干预展示”的端到端闭环。核心业务逻辑、图像上传全链路与推荐策略均可运行，满足课程与竞赛场景下的系统演示和持续迭代需求。后续将进一步补充自动化测试与容器化性能压测，提升工程化成熟度。'
]
for line in sec6_lines:
    p = insert_after(p, line)

doc.save(out_path)
print("SAVED:", out_path)

