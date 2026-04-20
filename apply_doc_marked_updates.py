from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SRC = r"c:\Users\21122\xwechat_files\wxid_p6jnc01rxek322_9c5c\msg\file\2026-04\chapter1_9_real_rewrite(2).docx"
OUT = r"c:\Users\21122\Desktop\HealthSystem\chapter1_9_real_rewrite_marked_fixed.docx"


def find_first(paragraphs, text: str):
    for i, p in enumerate(paragraphs):
        if p.text.strip() == text:
            return i
    return None


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def insert_after(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    out = Paragraph(new_p, paragraph._parent)
    out.add_run(text)
    return out


def main():
    doc = Document(SRC)
    ps = doc.paragraphs

    # 1) 处理 1.4 修改点
    mod_14 = None
    for i, p in enumerate(ps):
        t = p.text.strip()
        if t.startswith("（修改点：不要做项目当前完成度说明"):
            mod_14 = i
            break
    if mod_14 is not None:
        delete_paragraph(doc.paragraphs[mod_14])

    ps = doc.paragraphs
    idx_14 = find_first(ps, "1.4 项目当前完成度说明")
    if idx_14 is not None:
        doc.paragraphs[idx_14].text = "1.4 作品整体架构"
        # 替换下一段正文
        if idx_14 + 1 < len(doc.paragraphs):
            doc.paragraphs[idx_14 + 1].text = (
                "本作品采用“数据层→模型层→协同决策层→服务层→应用层”的全链路架构。"
                "数据层统一管理结构化健康数据、病种影像数据、文章与医院基础数据；模型层包含三病结构化模型与三病图像模型共6个模型；"
                "协同决策层负责结构化与图像结果融合（最终概率按6:4加权）及风险分层；服务层由FastAPI提供鉴权、预测、推荐、图像文件与管理接口；"
                "应用层由用户端与管理端页面承载，输出风险评估、个性化健康指南推荐、就近医院推荐和干预方案展示。"
                "当前版本支持本地演示部署，同时保留向多节点部署扩展的工程边界。"
            )

    # 2) 处理第7章修改点与重写
    ps = doc.paragraphs
    idx_ch7 = find_first(ps, "第7章 作品特色与创新点")
    idx_ch8 = find_first(ps, "第8章 应用推广与价值分析")
    if idx_ch7 is not None and idx_ch8 is not None and idx_ch8 > idx_ch7:
        # 删除第7章到第8章之间的全部旧内容（保留第7章标题）
        for i in range(idx_ch8 - 1, idx_ch7, -1):
            delete_paragraph(doc.paragraphs[i])

        # 插入新版第7章内容
        p = doc.paragraphs[idx_ch7]
        lines = [
            "7.1 多模态数据协同创新",
            "系统将结构化健康数据与医学影像数据统一到同一风险评估链路中，针对肝病、糖尿病、脑卒中分别构建结构化模型与图像模型，形成“三病×双模态”的协同框架。相比单一数据源方案，该设计更适应真实场景中“信息不完整、来源异构”的健康管理需求。",
            "7.2 融合决策与风险分层创新",
            "在预测阶段引入结构化与图像联合决策机制：双路数据同时可用时按6:4加权输出最终概率，单路缺失时自动退化为可用路径，保证结果连续可用。风险分层采用<30%低风险、30%-60%中风险、>=60%高风险，便于医生与用户快速理解和执行后续管理动作。",
            "7.3 个性化推荐与及时就医引导创新",
            "健康指南推荐不再采用泛化推送，而是基于病种与风险等级做规则化精准匹配；在中高风险场景中优先前置病种关键词相关内容，提升建议针对性。医院推荐结合定位与距离排序，为用户提供“及时就医、就近就医”的直接路径，强化从风险识别到行动执行的闭环。",
            "7.4 用户交互体验创新",
            "前端重点优化了关键使用环节：文章卡片支持全文弹窗阅读、正文自动分段与配图替换；影像上传支持已上传内容回显、覆盖修改与删除重传；风险页、干预页与数据采集页之间实现状态联动，显著降低用户重复操作成本，提升流程可用性。",
            "7.5 可视化与可解释性创新",
            "系统输出不仅包含概率与等级，还提供影响因素展示与融合明细，帮助用户理解“为什么是这个结果”。在页面层通过分病种风险标签、因素展示与推荐依据提示，实现“可看懂、可追溯、可执行”的可视化表达，提升系统可信度。",
            "7.6 工程可维护性创新",
            "后端通过模块化服务、统一字段映射和脚本化数据处理保障迭代效率；前端通过组件化改造和回归修复机制持续提升稳定性。项目已完成多轮问题修复（如完成度误判、上传白屏、图像回显缺失），体现了从功能实现到工程质量提升的持续演进能力。",
        ]
        for line in lines:
            p = insert_after(p, line)

    # 3) 删除第7章处的“修改点”残留（若仍存在）
    for i in range(len(doc.paragraphs) - 1, -1, -1):
        t = doc.paragraphs[i].text.strip()
        if t.startswith("（修改点：修改第7章"):
            delete_paragraph(doc.paragraphs[i])

    doc.save(OUT)
    print("SAVED:", OUT)


if __name__ == "__main__":
    main()
