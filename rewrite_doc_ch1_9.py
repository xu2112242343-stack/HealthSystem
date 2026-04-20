from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import os


BASE_DIR = r"c:\Users\21122\xwechat_files\wxid_p6jnc01rxek322_9c5c\msg\file\2026-04"
OUT_PATH = r"c:\Users\21122\Desktop\HealthSystem\chapter1_9_real_rewrite.docx"


def pick_target_docx(base_dir: str) -> str:
    files = [f for f in os.listdir(base_dir) if f.lower().endswith(".docx")]
    # 过滤明显不是目标的历史文件
    files = [f for f in files if "ex02" not in f.lower()]
    if not files:
        raise RuntimeError("未找到目标 docx 文件。")
    # 选最新修改的 docx
    files.sort(key=lambda x: os.path.getmtime(os.path.join(base_dir, x)), reverse=True)
    return os.path.join(base_dir, files[0])


def find_idx(paragraphs, startswith: str):
    for i, p in enumerate(paragraphs):
        if p.text.strip().startswith(startswith):
            return i
    return None


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def insert_after(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    out = Paragraph(new_p, paragraph._parent)
    out.add_run(text)
    return out


CHAPTER_LINES = {
    "第1章": [
        "1.1 研究背景与意义",
        "肝病、糖尿病与脑卒中具有“患病基数大、共病风险高、早筛窗口关键”的特点。现有健康管理系统常见问题是：仅做问卷评分、缺少图像协同、缺少可持续干预闭环。本项目围绕三病协同风险评估与干预，构建了“数据采集-风险预测-可解释展示-个性化管理”的完整流程。",
        "1.2 作品核心目标",
        "（1）实现三病风险评估：基于结构化健康数据完成肝病、糖尿病、脑卒中的概率预测与风险分层。",
        "（2）实现图像协同能力：支持三病对应影像上传、落盘、回显与删除，并将图像概率并入最终风险计算。",
        "（3）实现健康管理闭环：在风险结果基础上提供个性化健康指南推荐、医院就近推荐和干预方案生成入口。",
        "1.3 作品核心功能",
        "用户端：问卷填写、风险评估、干预方案查看、健康指南全文弹窗、医院按距离推荐、病种影像上传与管理。",
        "管理端：用户与医生账户管理、健康指南维护、数据查看与管理。",
        "算法侧：三病风险预测、风险等级映射、影响因子展示、结构化与图像加权融合（6:4）。",
        "1.4 项目当前完成度说明",
        "本项目已完成前后端主链路功能开发和关键问题修复，可用于课程展示与竞赛演示。当前部署形态以本地开发/演示为主，未宣称已完成生产级多节点集群落地。"
    ],
    "第2章": [
        "2.1 问题定义",
        "项目聚焦三个核心问题：第一，如何在用户可获取的数据条件下实现三病风险评估；第二，如何把图像数据纳入已有结构化预测链路；第三，如何把预测结果转化为可执行的健康管理建议。",
        "2.2 需求分析",
        "功能需求：用户注册登录、问卷保存、风险预测、影像上传管理、文章推荐、医院推荐、干预方案展示。",
        "数据需求：统一 user_info 字段体系，支持问卷字段映射、CSV 导入、图像路径字段持久化。",
        "算法需求：支持结构化模型优先可用，并在图像存在时进行融合；缺失任一路数据时系统可稳定退化。",
        "工程需求：接口可复用、前端交互可回显、关键页面无白屏、核心改动可回归验证。",
        "2.3 约束与边界",
        "本项目属于健康管理辅助系统，不提供临床诊断结论；所有预测结果用于风险提示与管理参考。大模型仅用于干预文本生成，不替代医生判断。",
        "2.4 解决思路",
        "采用“规则优先 + 可解释优先 + 可运行优先”的实现路径：先保证全链路可跑通，再逐步增强模型融合、推荐精细化与交互体验。"
    ],
    "第3章": [
        "3.1 数据来源与数据组织",
        "项目数据包含三类：用户结构化健康数据（问卷与指标）、病种文章数据、医院地理数据；另支持外部 CSV 转换导入用于测试账号构建。",
        "3.2 结构化数据处理",
        "后端通过 user_info 统一承载用户健康字段，questionnaire_save 模块将数据库记录映射为预测输入字典；并维护衍生指标（如部分比值）以支持模型与解释展示。",
        "3.3 图像数据处理",
        "在 user_info 表新增 liver_image_path、diabetes_image_path、stroke_image_path 三个字段。上传时后端保存文件并写入路径；读取时通过元数据接口与文件接口返回；删除时同步清理路径与磁盘文件。",
        "3.4 外部数据接入与批量构建",
        "项目提供 CSV 字段重命名与批量导入脚本，可把外部数据转换到 user_info 字段体系，并按规则批量创建测试账号，便于复现实验与演示。",
        "3.5 数据质量与一致性",
        "系统在输入映射、字段类型转换、缺失值兜底和接口异常处理方面做了工程保护，确保不完整数据下仍可返回可解释结果。"
    ],
    "第4章": [
        "4.1 技术架构",
        "后端采用 FastAPI + SQLAlchemy + MySQL，前端采用 React + TypeScript + Tailwind CSS。通过 REST API 完成用户数据、预测服务与干预服务的连接。",
        "4.2 三病风险预测流程",
        "核心入口为 risk_engine.predict_triple：先计算结构化概率与因子，再根据可用性选择模型输出或启发式输出，最终生成概率、分数、风险等级和解释因子。",
        "4.3 风险分层与评分",
        "按当前规则：概率 <30% 为低风险，30%-60% 为中风险，>=60% 为高风险。页面展示包含三病分值、等级标签与综合风险视图。",
        "4.4 结构化与图像融合策略",
        "当结构化与图像概率同时存在时，最终概率按“结构化:图像=6:4”加权；仅单路存在时直接使用该路概率。该策略已接入三病预测返回结构。",
        "4.5 个性化推荐与干预",
        "健康指南推荐采用规则引擎：三病全低风险时推荐认知类与饮食/运动类；存在中高风险时按病种+风险等级筛选，并对标题关键词命中的文章前置。",
        "医院推荐基于定位与经纬度距离计算，按就近排序展示。",
        "4.6 可解释性与合规边界",
        "系统返回三病影响因素列表、风险等级来源和融合明细；干预文本生成遵循“健康管理建议”边界，明确非诊断用途。"
    ],
    "第5章": [
        "5.1 系统工程架构与技术栈",
        "本项目采用前后端分离架构。前端为 React + TypeScript + Tailwind CSS（Vite 构建），后端为 FastAPI + SQLAlchemy + MySQL。核心能力围绕“肝病、糖尿病、脑卒中”三病风险评估与健康干预展开。",
        "后端以 REST API 提供用户、问卷、风险预测、健康指南、医院推荐、图像上传与读取等服务；前端实现用户端与管理端页面流程，完成数据采集、风险展示、推荐解释与干预交互。",
        "5.2 后端核心实现",
        "（1）统一预测入口：通过 risk_engine.predict_triple 计算三病风险，输出概率、风险等级、分数与可解释因素。",
        "（2）融合逻辑：三病最终概率支持结构化与图像 6:4 融合；缺一路时自动退化到单路输出。",
        "（3）推荐与医院能力：实现规则化健康指南推荐与按距离排序医院推荐。",
        "5.3 图像上传全链路实现",
        "user_info 新增三条图像路径字段，后端实现上传、元数据查询、文件读取、删除四类接口；前端实现回显、替换、删除与重传。",
        "5.4 前端核心页面实现",
        "干预页实现风险程度显示、指南卡片点击全文弹窗、定位与医院列表；数据采集页实现问卷与影像上传交互；风险页与首页实现分级结果展示与状态联动。",
        "5.5 关键工程修复",
        "已修复问卷完成度误判、影像上传按钮触发白屏、图像无法回显等问题，保证核心链路可连续操作。",
        "5.6 当前部署状态",
        "当前版本以本地开发/演示部署为主，已具备后续容器化与服务拆分基础。"
    ],
    "第6章": [
        "6.1 测试范围与方法",
        "测试覆盖模型预测链路、图像上传链路、推荐链路和关键页面交互。采用接口验证 + 页面流程回归 + 边界场景检查的组合方式。",
        "6.2 预测与融合验证",
        "验证结构化预测输出稳定；验证融合规则在三种场景下正确生效：双路都有、仅结构化、仅图像；验证30%与60%阈值边界分层正确。",
        "6.3 业务功能验证",
        "验证健康指南推荐规则（全低风险与中高风险两类逻辑）；验证文章弹窗排版与图片替换；验证医院定位与距离排序；验证风险页与干预页数据一致性。",
        "6.4 图像链路验证",
        "验证上传后路径写入、二次进入回显、覆盖上传更新、删除后清理四个关键动作；验证异常文件与失败状态提示。",
        "6.5 回归测试结果",
        "针对已修复问题进行回归：问卷完成度显示恢复正确，影像上传入口不再白屏，历史上传图像可查看并可修改/删除。",
        "6.6 测试结论",
        "当前版本已跑通“采集-预测-融合-推荐-干预”主闭环，功能可演示、逻辑可解释、可继续迭代。尚未进行系统化压测与多节点性能基准测试，后续将补充自动化测试与性能评估。"
    ],
    "第7章": [
        "7.1 多病种一体化风险链路",
        "项目将肝病、糖尿病、脑卒中三病放在统一流程中处理，减少用户重复录入与多系统切换成本。",
        "7.2 可落地的双路融合策略",
        "融合策略采用固定权重且支持缺路退化，工程可解释性强，便于在真实数据不完整场景中稳定运行。",
        "7.3 规则化个性推荐",
        "健康指南推荐从“泛推”改为“按病种与风险等级精准匹配”，并引入标题关键词前置，显著提升了结果相关性和可读性。",
        "7.4 图像管理闭环",
        "实现了上传、落盘、入库、回显、修改、删除的完整闭环，解决了传统演示系统“只能上传不能管理”的问题。",
        "7.5 工程迭代能力",
        "通过脚本化数据处理、接口分层和前端组件化改造，项目在持续修复与功能扩展方面具备较好可维护性。"
    ],
    "第8章": [
        "8.1 应用场景",
        "适用于高校健康管理教学演示、社区慢病风险初筛、体检后个体化健康指导等场景。",
        "8.2 应用价值",
        "对用户：获得直观风险结果、重点因素提示和可执行干预建议。",
        "对基层机构：可作为低门槛数字化辅助工具，提升初筛与随访效率。",
        "对团队研发：形成了从模型到产品化页面的端到端实践样板。",
        "8.3 推广可行性",
        "项目技术栈通用、部署门槛较低，支持逐步扩展为容器化部署和更完整的运维方案。",
        "8.4 已知限制",
        "当前以本地演示部署为主；部分模型依赖本地权重与环境；系统尚未完成全面自动化测试与标准化外部评测。"
    ],
    "第9章": [
        "9.1 工作总结",
        "本项目已完成三病风险评估、图像数据接入、6:4 融合计算、个性化指南推荐、医院距离推荐和干预展示等核心能力，形成可运行的健康管理闭环系统。",
        "9.2 主要收获",
        "在工程实践上，完成了数据字段统一、接口设计、前端交互闭环与关键缺陷修复；在业务逻辑上，完成了从“能跑”到“更个性化、更可解释”的迭代。",
        "9.3 后续计划",
        "（1）补充自动化测试与性能压测；（2）完善模型评估报告与可追踪版本管理；（3）推进容器化部署与服务拆分；（4）持续优化推荐与干预策略的医学一致性与用户体验。"
    ],
}


def rewrite_chapter(doc: Document, chapter_title: str, next_chapter_title: str):
    paragraphs = doc.paragraphs
    start = find_idx(paragraphs, chapter_title)
    end = find_idx(paragraphs, next_chapter_title)
    if start is None or end is None or end <= start:
        return False

    for i in range(end - 1, start, -1):
        delete_paragraph(doc.paragraphs[i])

    p = doc.paragraphs[start]
    for line in CHAPTER_LINES[chapter_title]:
        p = insert_after(p, line)
    return True


def main():
    target = pick_target_docx(BASE_DIR)
    doc = Document(target)

    ok = True
    ok &= rewrite_chapter(doc, "第1章", "第2章")
    ok &= rewrite_chapter(doc, "第2章", "第3章")
    ok &= rewrite_chapter(doc, "第3章", "第4章")
    ok &= rewrite_chapter(doc, "第4章", "第5章")
    ok &= rewrite_chapter(doc, "第5章", "第6章")
    ok &= rewrite_chapter(doc, "第6章", "第7章")
    ok &= rewrite_chapter(doc, "第7章", "第8章")
    ok &= rewrite_chapter(doc, "第8章", "第9章")
    # 第9章写到“参考文献”前
    paragraphs = doc.paragraphs
    start = find_idx(paragraphs, "第9章")
    end = find_idx(paragraphs, "参考文献")
    if start is None or end is None or end <= start:
        ok = False
    else:
        for i in range(end - 1, start, -1):
            delete_paragraph(doc.paragraphs[i])
        p = doc.paragraphs[start]
        for line in CHAPTER_LINES["第9章"]:
            p = insert_after(p, line)

    if not ok:
        raise SystemExit("章节定位失败，未完整重写。")

    doc.save(OUT_PATH)
    print("SAVED:", OUT_PATH)
    print("SOURCE:", target)


if __name__ == "__main__":
    main()
