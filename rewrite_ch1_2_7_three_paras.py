from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SRC = r"c:\Users\21122\Desktop\HealthSystem\chapter1_9_real_rewrite_marked_fixed.docx"
OUT = r"c:\Users\21122\Desktop\HealthSystem\chapter1_2_7_threepoints.docx"


def find_first(paragraphs, text: str):
    for i, p in enumerate(paragraphs):
        if p.text.strip() == text:
            return i
    return None


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def insert_after(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    out = Paragraph(new_p, paragraph._parent)
    out.add_run(text)
    return out


def replace_chapter(doc: Document, chapter_title: str, next_chapter_title: str, lines: list[str]):
    start = find_first(doc.paragraphs, chapter_title)
    end = find_first(doc.paragraphs, next_chapter_title)
    if start is None or end is None or end <= start:
        raise RuntimeError(f"章节定位失败: {chapter_title}")
    for i in range(end - 1, start, -1):
        delete_paragraph(doc.paragraphs[i])
    p = doc.paragraphs[start]
    for line in lines:
        p = insert_after(p, line)


def main():
    doc = Document(SRC)

    ch1_lines = [
        "1.1 研究背景与现实意义",
        "本作品聚焦肝病、糖尿病、脑卒中三类慢病在早筛与持续管理中的共性问题，整体思路是把分散的问卷指标、影像信息和干预服务整合到同一系统链路中，以降低用户多平台切换与重复录入成本，并提升从风险识别到后续管理的执行效率，使系统既能用于课程与竞赛展示，也能贴近真实健康管理场景。",
        "1.2 作品整体目标",
        "作品整体目标分为实现三病风险评估、实现图像协同能力与实现健康管理闭环三部分。实现三病风险评估基于结构化健康数据完成肝病、糖尿病、脑卒中的概率预测与风险分层。实现图像协同能力支持三病对应影像上传、落盘、回显与删除，并将图像概率并入最终风险计算。实现健康管理闭环是在风险结果基础上提供个性化健康指南推荐、医院就近推荐和干预方案展示入口。",
        "1.3 作品整体架构",
        "系统采用“数据层—模型层—协同决策层—服务层—应用层”的端到端架构，数据层统一管理用户健康数据、影像路径与推荐基础数据，模型层包含三病结构化与三病图像共6个模型，协同决策层负责结构化与图像结果融合及分级映射，服务层由后端接口承接预测与推荐能力，应用层由用户端和管理端页面承载完整业务流程。"
    ]

    ch2_lines = [
        "2.1 核心问题界定",
        "第一个核心问题是如何在用户可获得的数据条件下稳定完成三病风险评估，第二个核心问题是如何把影像能力接入既有结构化预测链路并保证结果可解释，第三个核心问题是如何让预测结果转化为用户可执行的后续动作，从而避免系统停留在“只给分数、不指导行为”的阶段。",
        "2.2 关键需求归纳",
        "整体需求可归纳为三类：功能需求上要完成采集、预测、上传、推荐与管理的主链路闭环，数据需求上要完成 user_info 字段统一、外部数据映射和图像路径持久化，工程需求上要保证接口可复用、页面可回显、关键异常可兜底并支持持续回归修复，以确保系统在迭代中保持可用性和一致性。",
        "2.3 约束条件与解决路径",
        "本作品定位为健康管理辅助系统而非临床诊断系统，输出用于风险提示与管理建议，因此在实现上采用“规则优先、可解释优先、可运行优先”的路径，先保证主流程可跑通，再逐步增强融合策略、推荐精度和交互体验，同时通过明确边界控制内容合规性和工程稳定性。"
    ]

    ch7_lines = [
        "7.1 多模态协同与融合策略创新",
        "本作品的核心创新之一是把结构化健康数据与医学影像数据纳入统一评估链路，并针对三病构建“结构化+图像”的双通道协同机制，在双路可用时采用固定权重融合、单路缺失时自动退化输出，既保证了结果连续可用，也提高了异构数据场景下的工程实用性与可解释性。",
        "7.2 个性化推荐与及时就医引导创新",
        "系统将风险结果直接映射到后续管理动作，形成“风险识别—内容推荐—就医引导—干预展示”的闭环，其中健康指南按病种和风险等级进行规则化精准匹配并对高相关内容前置，医院推荐结合定位与距离排序提供就近就医路径，使用户能从评估结果快速进入可执行的健康管理与就医决策。",
        "7.3 交互可视化与工程可维护性创新",
        "在用户体验与工程侧，作品通过风险可视化展示、文章全文弹窗、影像上传回显与修改删除等高频交互优化降低使用门槛，并通过字段映射统一、模块化接口与问题回归修复机制提升系统可维护性，使项目在功能完整度、可视化表达和持续迭代能力上形成可落地的综合优势。"
    ]

    replace_chapter(doc, "第1章 作品概述", "第2章 问题与需求分析", ch1_lines)
    replace_chapter(doc, "第2章 问题与需求分析", "第3章 数据集构建与数据预处理", ch2_lines)
    replace_chapter(doc, "第7章 作品特色与创新点", "第8章 应用推广与价值分析", ch7_lines)

    doc.save(OUT)
    print("SAVED:", OUT)


if __name__ == "__main__":
    main()
